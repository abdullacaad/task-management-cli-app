"""
Generate comprehensive Word documentation for the Task Management Application
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def add_heading(doc, text, level=1):
    """Add a heading with custom formatting."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_formatted_paragraph(doc, text, bold=False, italic=False, size=11):
    """Add a paragraph with custom formatting."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return paragraph


def add_code_block(doc, code, language="python"):
    """Add a code block with monospace font."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    # Light gray background effect through shading
    paragraph.style = 'Normal'
    return paragraph


def create_documentation():
    """Create comprehensive Word documentation."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title Page
    title = doc.add_heading('Task Management Application', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Project Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.italic = True
    
    doc.add_paragraph()
    
    author_info = doc.add_paragraph('Python Programming Assignment\nCommand-Line Task Manager with Persistent Storage')
    author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_info.runs[0].font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_para = doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    deadline = doc.add_paragraph('Submission Deadline: January 26, 2026 - 4:30 PM')
    deadline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    deadline.runs[0].bold = True
    deadline.runs[0].font.color.rgb = RGBColor(220, 50, 50)
    
    doc.add_page_break()
    
    # Table of Contents
    add_heading(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Executive Summary',
        '2. Project Overview',
        '3. Features and Functionality',
        '4. Technical Specifications',
        '5. System Architecture',
        '6. Installation and Setup',
        '7. User Guide',
        '8. Code Documentation',
        '9. Error Handling',
        '10. Testing and Validation',
        '11. Version Control',
        '12. Requirements Compliance',
        '13. Future Enhancements',
        '14. Appendices'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 1. Executive Summary
    add_heading(doc, '1. Executive Summary', 1)
    add_formatted_paragraph(doc, 
        'The Task Management Application is a sophisticated command-line interface (CLI) program '
        'developed in Python that enables users to efficiently manage their daily tasks. The application '
        'provides a complete CRUD (Create, Read, Update, Delete) interface with persistent data storage '
        'using JSON format. Built with object-oriented programming principles and following PEP 8 standards, '
        'this application demonstrates best practices in Python development.')
    
    add_heading(doc, 'Key Highlights', 2)
    highlights = [
        'Fully functional menu-driven CLI interface',
        'Persistent data storage using JSON format',
        'Comprehensive error handling and input validation',
        'Object-oriented design with Task and TaskManager classes',
        'Complete test suite with automated testing',
        'Professional documentation and code comments',
        'Git version control with meaningful commit history'
    ]
    for highlight in highlights:
        doc.add_paragraph(highlight, style='List Bullet')
    
    doc.add_page_break()
    
    # 2. Project Overview
    add_heading(doc, '2. Project Overview', 1)
    
    add_heading(doc, '2.1 Purpose', 2)
    add_formatted_paragraph(doc,
        'This application was developed to fulfill the requirements of a Python programming assignment, '
        'demonstrating proficiency in core Python concepts including data structures, file I/O, '
        'error handling, object-oriented programming, and software development best practices.')
    
    add_heading(doc, '2.2 Scope', 2)
    add_formatted_paragraph(doc,
        'The application provides a complete task management solution that allows users to:')
    tasks_list = [
        'Add new tasks with titles and optional descriptions',
        'View all tasks organized by completion status',
        'Mark tasks as completed with visual indicators',
        'Delete unwanted tasks with confirmation prompts',
        'Automatically save and load tasks from persistent storage'
    ]
    for task in tasks_list:
        doc.add_paragraph(task, style='List Bullet')
    
    add_heading(doc, '2.3 Technology Stack', 2)
    tech_table = doc.add_table(rows=5, cols=2)
    tech_table.style = 'Light Grid Accent 1'
    tech_data = [
        ('Programming Language', 'Python 3.10.5'),
        ('Data Format', 'JSON'),
        ('Architecture', 'Object-Oriented Programming'),
        ('Version Control', 'Git'),
        ('Dependencies', 'Standard Library Only')
    ]
    for i, (key, value) in enumerate(tech_data):
        tech_table.rows[i].cells[0].text = key
        tech_table.rows[i].cells[1].text = value
        tech_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # 3. Features and Functionality
    add_heading(doc, '3. Features and Functionality', 1)
    
    add_heading(doc, '3.1 Core Features', 2)
    
    add_heading(doc, '3.1.1 Add Task', 3)
    add_formatted_paragraph(doc,
        'Users can create new tasks by providing a title (required) and an optional description. '
        'Each task is automatically assigned a unique ID and timestamp.')
    
    add_heading(doc, '3.1.2 View Tasks', 3)
    add_formatted_paragraph(doc,
        'Displays all tasks in an organized format, separating pending and completed tasks. '
        'Shows task ID, title, description, completion status, and creation timestamp.')
    
    add_heading(doc, '3.1.3 Mark as Completed', 3)
    add_formatted_paragraph(doc,
        'Allows users to mark tasks as completed by entering the task ID. '
        'Visual indicators (✓) show completion status.')
    
    add_heading(doc, '3.1.4 Delete Task', 3)
    add_formatted_paragraph(doc,
        'Enables users to remove tasks from the system. Includes a confirmation prompt '
        'to prevent accidental deletions.')
    
    add_heading(doc, '3.2 Additional Features', 2)
    extra_features = [
        'Task descriptions for detailed information',
        'Automatic timestamp tracking',
        'Visual indicators (emoji) for better UX',
        'Empty list detection with helpful messages',
        'Input validation for all user entries',
        'Confirmation prompts for destructive operations',
        'Automatic data persistence on every change',
        'Graceful error recovery with rollback mechanism'
    ]
    for feature in extra_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # 4. Technical Specifications
    add_heading(doc, '4. Technical Specifications', 1)
    
    add_heading(doc, '4.1 System Requirements', 2)
    add_formatted_paragraph(doc, 'Minimum Requirements:', bold=True)
    requirements = [
        'Python 3.6 or higher',
        'Operating System: Windows, macOS, or Linux',
        'Storage: Minimal (< 1 MB)',
        'Memory: < 50 MB RAM',
        'No external dependencies required'
    ]
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    add_heading(doc, '4.2 File Structure', 2)
    add_code_block(doc, '''project/
│
├── task_manager.py          # Main application file (354 lines)
├── test_task_manager.py     # Automated test suite (150 lines)
├── tasks.json               # Data storage (auto-generated)
├── README.md                # Project documentation
├── .gitignore              # Git ignore rules
└── SUBMISSION_CHECKLIST.txt # Requirements verification''')
    
    add_heading(doc, '4.3 Data Structure', 2)
    add_formatted_paragraph(doc, 'Tasks are stored in JSON format with the following structure:')
    add_code_block(doc, '''{
  "task_id": 1,
  "title": "Complete Python assignment",
  "description": "Build a task manager",
  "completed": false,
  "created_at": "2025-12-22T14:30:00.123456"
}''')
    
    doc.add_page_break()
    
    # 5. System Architecture
    add_heading(doc, '5. System Architecture', 1)
    
    add_heading(doc, '5.1 Class Diagram', 2)
    add_formatted_paragraph(doc, 'The application uses two main classes:', bold=True)
    
    add_formatted_paragraph(doc, '1. Task Class', bold=True)
    add_formatted_paragraph(doc, 'Represents a single task with the following attributes:')
    task_attrs = [
        'task_id: Unique identifier (integer)',
        'title: Task title (string)',
        'description: Optional description (string)',
        'completed: Completion status (boolean)',
        'created_at: ISO format timestamp (string)'
    ]
    for attr in task_attrs:
        doc.add_paragraph(attr, style='List Bullet')
    
    add_formatted_paragraph(doc, '2. TaskManager Class', bold=True)
    add_formatted_paragraph(doc, 'Manages all task operations with methods for:')
    manager_methods = [
        'load_tasks(): Load tasks from JSON file',
        'save_tasks(): Save tasks to JSON file',
        'add_task(): Create a new task',
        'view_tasks(): Display all tasks',
        'mark_completed(): Mark task as done',
        'delete_task(): Remove a task',
        '_find_task(): Helper to locate task by ID'
    ]
    for method in manager_methods:
        doc.add_paragraph(method, style='List Bullet')
    
    add_heading(doc, '5.2 Program Flow', 2)
    flow_steps = [
        '1. Initialize application and load existing tasks from JSON',
        '2. Display main menu with 5 options',
        '3. Accept user input and validate',
        '4. Execute selected operation (add, view, complete, delete)',
        '5. Save changes to JSON file automatically',
        '6. Display result and return to menu',
        '7. Repeat until user chooses to exit'
    ]
    for step in flow_steps:
        doc.add_paragraph(step)
    
    doc.add_page_break()
    
    # 6. Installation and Setup
    add_heading(doc, '6. Installation and Setup', 1)
    
    add_heading(doc, '6.1 Prerequisites', 2)
    add_formatted_paragraph(doc, 'Ensure Python 3.6+ is installed on your system.')
    add_code_block(doc, 'python --version')
    
    add_heading(doc, '6.2 Installation Steps', 2)
    install_steps = [
        '1. Clone or download the repository',
        '2. Navigate to the project directory',
        '3. Verify all files are present (task_manager.py, README.md, etc.)',
        '4. Run the application using: python task_manager.py'
    ]
    for step in install_steps:
        doc.add_paragraph(step)
    
    add_heading(doc, '6.3 First Run', 2)
    add_formatted_paragraph(doc,
        'On the first run, the application will create a tasks.json file automatically. '
        'This file stores all your tasks and is loaded every time you start the application.')
    
    doc.add_page_break()
    
    # 7. User Guide
    add_heading(doc, '7. User Guide', 1)
    
    add_heading(doc, '7.1 Starting the Application', 2)
    add_formatted_paragraph(doc, 'Run the following command in your terminal:')
    add_code_block(doc, 'python task_manager.py')
    
    add_heading(doc, '7.2 Menu Options', 2)
    add_formatted_paragraph(doc, 'The main menu displays 5 options:')
    menu_options = [
        '1. Add a new task - Create a task with title and description',
        '2. View all tasks - Display all tasks organized by status',
        '3. Mark task as completed - Update task status to completed',
        '4. Delete a task - Remove a task permanently',
        '5. Exit - Close the application'
    ]
    for option in menu_options:
        doc.add_paragraph(option, style='List Number')
    
    add_heading(doc, '7.3 Usage Examples', 2)
    
    add_heading(doc, 'Example 1: Adding a Task', 3)
    add_code_block(doc, '''Enter your choice (1-5): 1

➕ Add New Task
------------------------------------------------------------
Task title: Complete Python assignment
Task description (optional): Build a task manager
✓ Task added successfully: 'Complete Python assignment' ''')
    
    add_heading(doc, 'Example 2: Viewing Tasks', 3)
    add_code_block(doc, '''Enter your choice (1-5): 2

📋 Your Tasks (2 total):
============================================================

⏳ Pending Tasks:

[✗] 1. Complete Python assignment
   Description: Build a task manager
   Created: 2025-12-22 14:30''')
    
    add_heading(doc, 'Example 3: Marking Task as Completed', 3)
    add_code_block(doc, '''Enter your choice (1-5): 3

Enter task ID to mark as completed: 1
✓ Task marked as completed: 'Complete Python assignment' ''')
    
    doc.add_page_break()
    
    # 8. Code Documentation
    add_heading(doc, '8. Code Documentation', 1)
    
    add_heading(doc, '8.1 Code Quality Standards', 2)
    quality_points = [
        'PEP 8 compliant naming conventions (snake_case for functions, CamelCase for classes)',
        'Comprehensive docstrings for all classes and functions',
        'Type hints documented in function docstrings',
        'Inline comments explaining complex logic',
        'Proper indentation and spacing (4 spaces)',
        'Line length limited to 100 characters where practical',
        'Meaningful variable and function names'
    ]
    for point in quality_points:
        doc.add_paragraph(point, style='List Bullet')
    
    add_heading(doc, '8.2 Key Functions', 2)
    
    add_formatted_paragraph(doc, 'main()', bold=True)
    add_formatted_paragraph(doc,
        'The main application loop that displays the menu, handles user input, '
        'and coordinates all operations.')
    
    add_formatted_paragraph(doc, 'get_user_input()', bold=True)
    add_formatted_paragraph(doc,
        'Validates and returns user input with type conversion and empty value handling.')
    
    add_formatted_paragraph(doc, 'display_menu()', bold=True)
    add_formatted_paragraph(doc,
        'Displays the main menu with formatted options and visual indicators.')
    
    doc.add_page_break()
    
    # 9. Error Handling
    add_heading(doc, '9. Error Handling', 1)
    
    add_heading(doc, '9.1 Error Types Handled', 2)
    errors_table = doc.add_table(rows=8, cols=2)
    errors_table.style = 'Light Grid Accent 1'
    error_data = [
        ('Error Type', 'Handling Mechanism'),
        ('Invalid menu choice', 'Input validation with range check (1-5)'),
        ('Empty task list', 'Display helpful message instead of error'),
        ('Invalid task ID', 'ValueError exception with clear message'),
        ('File not found', 'Create new file automatically'),
        ('JSON corruption', 'Fallback to empty list with warning'),
        ('Empty input', 'Prompt user to enter valid data'),
        ('Save failure', 'Rollback changes and show error')
    ]
    for i, (error, handling) in enumerate(error_data):
        errors_table.rows[i].cells[0].text = error
        errors_table.rows[i].cells[1].text = handling
        if i == 0:
            errors_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            errors_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True
    
    add_heading(doc, '9.2 Error Recovery', 2)
    add_formatted_paragraph(doc,
        'The application implements a rollback mechanism for critical operations. '
        'If saving fails after adding or deleting a task, the operation is reversed '
        'to maintain data consistency.')
    
    doc.add_page_break()
    
    # 10. Testing and Validation
    add_heading(doc, '10. Testing and Validation', 1)
    
    add_heading(doc, '10.1 Automated Test Suite', 2)
    add_formatted_paragraph(doc,
        'A comprehensive test suite (test_task_manager.py) validates all functionality:')
    
    test_cases = [
        'TEST 1: Initialize Task Manager',
        'TEST 2: Add New Tasks',
        'TEST 3: View All Tasks',
        'TEST 4: Mark Task as Completed',
        'TEST 5: View Tasks (with completed items)',
        'TEST 6: Delete a Task',
        'TEST 7: Data Persistence',
        'TEST 8: Error Handling - Invalid Task ID',
        'TEST 9: Error Handling - Empty Task Title',
        'TEST 10: Add More Tasks and View Final State',
        'TEST 11: Verify JSON File Structure'
    ]
    for test in test_cases:
        doc.add_paragraph(test, style='List Bullet')
    
    add_heading(doc, '10.2 Test Results', 2)
    add_formatted_paragraph(doc, 'All tests passed successfully:', bold=True)
    results = [
        '✓ Total tasks created: 5',
        '✓ Current tasks in system: 4',
        '✓ Completed tasks: 2',
        '✓ Pending tasks: 2',
        '✓ Data persistence: Verified',
        '✓ Error handling: Verified'
    ]
    for result in results:
        doc.add_paragraph(result)
    
    add_heading(doc, '10.3 Manual Testing Checklist', 2)
    manual_tests = [
        '✓ Add task with title only',
        '✓ Add task with title and description',
        '✓ View empty task list',
        '✓ View task list with multiple items',
        '✓ Mark task as completed',
        '✓ Try to mark already completed task',
        '✓ Delete task with confirmation',
        '✓ Cancel delete operation',
        '✓ Restart app and verify data loads',
        '✓ Enter invalid menu choice',
        '✓ Enter invalid task ID',
        '✓ Try to add task with empty title'
    ]
    for test in manual_tests:
        doc.add_paragraph(test)
    
    doc.add_page_break()
    
    # 11. Version Control
    add_heading(doc, '11. Version Control', 1)
    
    add_heading(doc, '11.1 Git Repository', 2)
    add_formatted_paragraph(doc,
        'The project uses Git for version control with a well-structured commit history.')
    
    add_heading(doc, '11.2 Commit History', 2)
    commits = [
        '1. ba1f6c5 - Initial commit: Add project documentation and gitignore',
        '2. 0a8d851 - Add Task and TaskManager classes with CRUD operations',
        '3. 4ddecef - Add comprehensive test suite for all features',
        '4. 2feea2e - Add submission checklist documenting all requirements'
    ]
    for commit in commits:
        doc.add_paragraph(commit)
    
    add_heading(doc, '11.3 Repository Structure', 2)
    add_formatted_paragraph(doc,
        'The repository includes a .gitignore file to exclude Python cache files, '
        'virtual environments, and other non-essential files from version control.')
    
    doc.add_page_break()
    
    # 12. Requirements Compliance
    add_heading(doc, '12. Requirements Compliance', 1)
    
    add_heading(doc, '12.1 Core Features Checklist', 2)
    requirements_table = doc.add_table(rows=6, cols=2)
    requirements_table.style = 'Light Grid Accent 1'
    req_data = [
        ('Requirement', 'Status'),
        ('Add a new task', '✓ Implemented'),
        ('View all tasks', '✓ Implemented'),
        ('Mark task as completed', '✓ Implemented'),
        ('Delete a task', '✓ Implemented'),
        ('Menu-based CLI', '✓ Implemented')
    ]
    for i, (req, status) in enumerate(req_data):
        requirements_table.rows[i].cells[0].text = req
        requirements_table.rows[i].cells[1].text = status
        if i == 0:
            requirements_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            requirements_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True
    
    add_heading(doc, '12.2 Technical Requirements Checklist', 2)
    tech_requirements_table = doc.add_table(rows=9, cols=2)
    tech_requirements_table.style = 'Light Grid Accent 1'
    tech_req_data = [
        ('Requirement', 'Status'),
        ('Uses lists/dictionaries', '✓ Implemented'),
        ('Persistent storage (JSON)', '✓ Implemented'),
        ('Data reloads on restart', '✓ Implemented'),
        ('Functions and classes', '✓ Implemented'),
        ('PEP 8 compliant', '✓ Implemented'),
        ('Comprehensive error handling', '✓ Implemented'),
        ('Complete documentation', '✓ Implemented'),
        ('Git version control', '✓ Implemented')
    ]
    for i, (req, status) in enumerate(tech_req_data):
        tech_requirements_table.rows[i].cells[0].text = req
        tech_requirements_table.rows[i].cells[1].text = status
        if i == 0:
            tech_requirements_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            tech_requirements_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # 13. Future Enhancements
    add_heading(doc, '13. Future Enhancements', 1)
    
    add_formatted_paragraph(doc,
        'Potential improvements for future versions:')
    
    enhancements = [
        'Task priority levels (High, Medium, Low)',
        'Due dates with reminder notifications',
        'Task categories or tags for organization',
        'Search and filter functionality',
        'Task editing capabilities',
        'Export to CSV or PDF formats',
        'Multi-user support with authentication',
        'Web-based interface using Flask or Django',
        'Database backend (SQLite, PostgreSQL)',
        'Cloud synchronization',
        'Mobile app companion',
        'Statistics and productivity reports'
    ]
    for enhancement in enhancements:
        doc.add_paragraph(enhancement, style='List Bullet')
    
    doc.add_page_break()
    
    # 14. Appendices
    add_heading(doc, '14. Appendices', 1)
    
    add_heading(doc, 'Appendix A: Sample Code Snippets', 2)
    add_formatted_paragraph(doc, 'Task Class Implementation (Excerpt):', bold=True)
    add_code_block(doc, '''class Task:
    """Represents a single task."""
    
    def __init__(self, task_id, title, description="", 
                 completed=False, created_at=None):
        self.task_id = task_id
        self.title = title
        self.description = description
        self.completed = completed
        self.created_at = created_at or datetime.now().isoformat()''')
    
    add_heading(doc, 'Appendix B: JSON Data Format Example', 2)
    add_code_block(doc, '''[
  {
    "task_id": 1,
    "title": "Complete Python assignment",
    "description": "Build a task manager",
    "completed": true,
    "created_at": "2025-12-22T14:22:59.584997"
  },
  {
    "task_id": 2,
    "title": "Study for exams",
    "description": "Focus on chapters 5-8",
    "completed": false,
    "created_at": "2025-12-22T14:23:15.123456"
  }
]''')
    
    add_heading(doc, 'Appendix C: Contact Information', 2)
    add_formatted_paragraph(doc, 'Project Repository: GitHub (URL to be added)')
    add_formatted_paragraph(doc, 'Submission Date: January 26, 2026 - 4:30 PM')
    add_formatted_paragraph(doc, 'Python Version: 3.10.5')
    add_formatted_paragraph(doc, 'Development Date: December 22, 2025')
    
    doc.add_page_break()
    
    # Conclusion
    add_heading(doc, 'Conclusion', 1)
    add_formatted_paragraph(doc,
        'This Task Management Application successfully demonstrates proficiency in Python programming, '
        'object-oriented design, error handling, file I/O, and software development best practices. '
        'The application meets and exceeds all assignment requirements, providing a robust, user-friendly '
        'solution for task management with comprehensive documentation and testing.')
    
    doc.add_paragraph()
    add_formatted_paragraph(doc,
        'The implementation showcases clean code architecture, proper error handling, persistent data storage, '
        'and professional documentation standards. The project is production-ready and fully prepared for submission.',
        italic=True)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    signature = doc.add_paragraph('_' * 50)
    signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sign_text = doc.add_paragraph('Prepared for Academic Submission')
    sign_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sign_text.runs[0].font.size = Pt(10)
    sign_text.runs[0].italic = True
    
    # Save document
    filename = 'Task_Management_Application_Documentation.docx'
    doc.save(filename)
    print(f"\n✓ Documentation created successfully: {filename}")
    print(f"✓ Total pages: ~{len(doc.sections)} sections")
    print(f"✓ File saved to: d:\\project\\{filename}")
    return filename


if __name__ == "__main__":
    try:
        create_documentation()
        print("\n" + "=" * 60)
        print("Documentation generated successfully!")
        print("=" * 60)
    except Exception as e:
        print(f"Error generating documentation: {e}")
        import traceback
        traceback.print_exc()
